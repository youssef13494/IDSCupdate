# ============================================
# FILE 1: rag.py
# ============================================
import os
from docx import Document as DocxDocument
from typing import List, Optional, Tuple
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_qdrant import QdrantVectorStore
from langchain_core.documents import Document
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams
from rank_bm25 import BM25Okapi
import numpy as np
import pickle
import re


class ArabicDocProcessor:
    _instance = None
    _lock = False
    
    def __new__(cls, persist_directory: str = "./vector_store"):
        if cls._instance is None:
            cls._instance = super(ArabicDocProcessor, cls).__new__(cls)
            cls._instance._initialized = False
        return cls._instance
    
    def __init__(self, persist_directory: str = "./vector_store"):
        if self._initialized:
            return
            
        self.persist_directory = persist_directory
        self.collection_name = "arabic_docs"
        
        self._embeddings = None
        self._client = None
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=300,
            separators=["\n\n", "\n", ".", "،", " ", ""]
        )
        self.chars_per_page = 3000
        
        self.bm25_index = None
        self.bm25_documents = []
        self.bm25_path = os.path.join(persist_directory, "bm25_index.pkl")
        
        self._initialized = True
    
    @property
    def embeddings(self):
        if self._embeddings is None:
            model_name = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
            self._embeddings = HuggingFaceEmbeddings(
                model_name=model_name,
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
        return self._embeddings
    
    @property
    def client(self):
        if self._client is None and not ArabicDocProcessor._lock:
            ArabicDocProcessor._lock = True
            try:
                self._client = QdrantClient(path=self.persist_directory)
                print("✅ Qdrant client initialized (file-based)")
            except Exception as e:
                print(f"❌ Could not use file-based Qdrant: {e}")
                print("Falling back to in-memory Qdrant client")
                self._client = QdrantClient(":memory:")
            finally:
                ArabicDocProcessor._lock = False
        return self._client
    
    def initialize(self):
        self._load_bm25()
        _ = self.client
        
    def _has_numbers(self, text: str) -> bool:
        return bool(re.search(r'\d', text))
    
    def _extract_numbers(self, text: str) -> List[str]:
        return re.findall(r'\d+(?:\.\d+)?', text)
    
    def _load_bm25(self):
        try:
            if os.path.exists(self.bm25_path):
                with open(self.bm25_path, 'rb') as f:
                    data = pickle.load(f)
                    self.bm25_index = data['bm25']
                    self.bm25_documents = data['documents']
                print(f"✅ Loaded BM25 index with {len(self.bm25_documents)} documents")
        except Exception as e:
            print(f"Could not load BM25 index: {str(e)}")
    
    def _save_bm25(self):
        try:
            os.makedirs(self.persist_directory, exist_ok=True)
            with open(self.bm25_path, 'wb') as f:
                pickle.dump({
                    'bm25': self.bm25_index,
                    'documents': self.bm25_documents
                }, f)
            print("✅ BM25 index saved")
        except Exception as e:
            print(f"Error saving BM25 index: {str(e)}")
    
    def _create_bm25_index(self, documents: List[Document]):
        self.bm25_documents = documents
        tokenized_corpus = []
        
        for doc in documents:
            text = doc.page_content
            text = re.sub(r'(\d+)', r' \1 ', text)
            tokens = text.split()
            tokenized_corpus.append(tokens)
        
        self.bm25_index = BM25Okapi(tokenized_corpus)
        self._save_bm25()
    
    def extract_text_from_docx(self, docx_path: str) -> Tuple[str, List[dict]]:
        paragraphs_with_pages = []
        current_char_count = 0
        
        try:
            doc = DocxDocument(docx_path)
            
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    estimated_page = (current_char_count // self.chars_per_page) + 1
                    paragraphs_with_pages.append({
                        'text': para_text,
                        'page': estimated_page
                    })
                    current_char_count += len(para_text)
            
            for table in doc.tables:
                table_text = ""
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            table_text += cell_text + " "
                
                if table_text:
                    estimated_page = (current_char_count // self.chars_per_page) + 1
                    paragraphs_with_pages.append({
                        'text': table_text,
                        'page': estimated_page
                    })
                    current_char_count += len(table_text)
                        
        except Exception as e:
            print(f"Error extracting text from {docx_path}: {str(e)}")
        
        full_text = "\n\n".join([p['text'] for p in paragraphs_with_pages])
        return full_text, paragraphs_with_pages
    
    def process_docx(self, docx_path: str, metadata: Optional[dict] = None) -> List[Document]:
        full_text, paragraphs_with_pages = self.extract_text_from_docx(docx_path)
        
        if not full_text.strip():
            return []
        
        if metadata is None:
            metadata = {}
        
        base_metadata = {
            "source": os.path.basename(docx_path),
            "type": "docx",
            "language": "arabic"
        }
        base_metadata.update(metadata)
        
        texts = self.text_splitter.split_text(full_text)
        
        documents = []
        for i, text_chunk in enumerate(texts):
            doc_metadata = base_metadata.copy()
            doc_metadata["chunk"] = i
            doc_metadata["has_numbers"] = self._has_numbers(text_chunk)
            doc_metadata["numbers"] = self._extract_numbers(text_chunk)
            
            chunk_position = full_text.find(text_chunk[:100])
            if chunk_position != -1:
                estimated_page = (chunk_position // self.chars_per_page) + 1
            else:
                estimated_page = (i * 1500 // self.chars_per_page) + 1
            
            doc_metadata["page"] = estimated_page
            documents.append(Document(page_content=text_chunk, metadata=doc_metadata))
    
        return documents
    
    def process_directory(self, directory_path: str) -> List[Document]:
        all_documents = []
        
        if not os.path.exists(directory_path):
            os.makedirs(directory_path, exist_ok=True)
            return all_documents
        
        for filename in os.listdir(directory_path):
            if filename.lower().endswith(('.docx', '.doc')):
                docx_path = os.path.join(directory_path, filename)
                print(f"Processing {filename}...")
                documents = self.process_docx(docx_path)
                all_documents.extend(documents)
        
        return all_documents
    
    def create_vector_store(self, documents: List[Document]):
        if not documents:
            print("No documents to process")
            return None
        
        sample_embedding = self.embeddings.embed_query("test")
        embedding_dim = len(sample_embedding)
        
        try:
            self.client.delete_collection(collection_name=self.collection_name)
        except:
            pass
        
        self.client.create_collection(
            collection_name=self.collection_name,
            vectors_config=VectorParams(size=embedding_dim, distance=Distance.COSINE)
        )
        
        vector_store = QdrantVectorStore(
            client=self.client,
            collection_name=self.collection_name,
            embedding=self.embeddings
        )
        
        vector_store.add_documents(documents)
        self._create_bm25_index(documents)
        
        print(f"✅ Vector store created with {len(documents)} documents")
        print(f"✅ BM25 index created")
        return vector_store
    
    def load_vector_store(self):
        try:
            collections = self.client.get_collections().collections
            if not any(c.name == self.collection_name for c in collections):
                return None
            
            vector_store = QdrantVectorStore(
                client=self.client,
                collection_name=self.collection_name,
                embedding=self.embeddings
            )
            return vector_store
        except Exception as e:
            print(f"Error loading vector store: {str(e)}")
            return None
    
    def _bm25_search(self, query: str, k: int = 80) -> List[Tuple[Document, float]]:
        if not self.bm25_index or not self.bm25_documents:
            return []
        
        query_processed = re.sub(r'(\d+)', r' \1 ', query)
        tokenized_query = query_processed.split()
        
        scores = self.bm25_index.get_scores(tokenized_query)
        
        query_numbers = set(self._extract_numbers(query))
        if query_numbers:
            for idx, doc in enumerate(self.bm25_documents):
                doc_numbers = set(doc.metadata.get('numbers', []))
                if query_numbers & doc_numbers:
                    scores[idx] *= 1.5
        
        top_indices = np.argsort(scores)[::-1][:k]
        
        results = []
        for idx in top_indices:
            if scores[idx] > 0:
                results.append((self.bm25_documents[idx], float(scores[idx])))
        
        return results
    
    def _dense_search(self, vector_store, query: str, k: int = 60) -> List[Tuple[Document, float]]:
        results = vector_store.similarity_search_with_score(query, k=k)
        return [(doc, float(score)) for doc, score in results]
    
    def _hybrid_search(self, vector_store, query: str, k: int = 40, alpha: float = 0.5) -> List[Document]:
        bm25_results = self._bm25_search(query, k=k*2)
        dense_results = self._dense_search(vector_store, query, k=k*2)
        
        def normalize_scores(results):
            if not results:
                return {}
            scores = [score for _, score in results]
            min_score = min(scores)
            max_score = max(scores)
            if max_score == min_score:
                return {i: 0.5 for i, _ in enumerate(results)}
            return {i: (score - min_score) / (max_score - min_score) 
                   for i, (_, score) in enumerate(results)}
        
        bm25_norm = normalize_scores(bm25_results)
        dense_norm = normalize_scores(dense_results)
        
        has_numbers = self._has_numbers(query)
        adjusted_alpha = 0.4 if has_numbers else alpha
        
        combined_scores = {}
        
        for i, (doc, _) in enumerate(bm25_results):
            doc_id = doc.page_content
            combined_scores[doc_id] = (1 - adjusted_alpha) * bm25_norm.get(i, 0)
        
        for i, (doc, _) in enumerate(dense_results):
            doc_id = doc.page_content
            if doc_id in combined_scores:
                combined_scores[doc_id] += adjusted_alpha * dense_norm.get(i, 0)
            else:
                combined_scores[doc_id] = adjusted_alpha * dense_norm.get(i, 0)
        
        doc_map = {}
        for doc, _ in bm25_results + dense_results:
            doc_map[doc.page_content] = doc
        
        sorted_docs = sorted(combined_scores.items(), key=lambda x: x[1], reverse=True)
        
        results = []
        seen = set()
        for doc_id, score in sorted_docs:
            if doc_id not in seen and doc_id in doc_map:
                results.append(doc_map[doc_id])
                seen.add(doc_id)
                if len(results) >= k:
                    break
        
        return results
    
    def search_documents(self, query: str, k: int = 40) -> List[Document]:
        vector_store = self.load_vector_store()
        if not vector_store:
            return []
        
        if self.bm25_index:
            print(f"✅ Using Hybrid Search (BM25 + Dense) - fetching top {k}")
            print(f"📊 Query has numbers: {self._has_numbers(query)}")
            return self._hybrid_search(vector_store, query, k=k, alpha=0.5)
        else:
            print(f"✅ Using Dense Search only - fetching top {k}")
            return vector_store.similarity_search(query, k=k)
    
    def get_collection_info(self) -> dict:
        try:
            collection_info = self.client.get_collection(self.collection_name)
            return {
                "count": collection_info.points_count,
                "name": self.collection_name,
                "has_bm25": self.bm25_index is not None,
                "chunk_size": 1500,
                "chunk_overlap": 300
            }
        except:
            return {
                "count": 0, 
                "name": self.collection_name,
                "has_bm25": False,
                "chunk_size": 1500,
                "chunk_overlap": 300
            }

arabic_processor = ArabicDocProcessor()
